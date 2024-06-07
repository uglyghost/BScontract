from tqdm import tqdm
from concurrent.futures import ThreadPoolExecutor, as_completed
from docx import Document
from aaa import revise_document
import process
import os
import re
import threading
from annotate import add_comments
from llm import get_correction_suggestion

url_base_list = [
    "https://api.rcouyi.com/v1/",
    "https://us.rcouyi.com/v1/",
    "https://us-1.rcouyi.com/v1/",
    "https://jp-5.rcouyi.com/v1/"
]


def process_file(file_data):
    for file in file_data:
        try:
            check_contract(file.get('file_path'), file.get('unique_filename'))
        except Exception as e:
            print("process_file捕获到异常：", str(e))
            continue


def check_contract(file_path, unique_filename):
    try:
        doc = Document(file_path)
        contract_comments, doc = check_text(doc)
        processed_filename = f"processed_{unique_filename}"
        revise_filename = f"processed_revise_{unique_filename}"
        processed_filepath = os.path.join('processed/', processed_filename)
        revise_filepath = os.path.join('processed/', revise_filename)
        doc.save(processed_filepath)
        add_comments(contract_comments, processed_filepath)
        print(contract_comments)
        for item in contract_comments:
            print(len(contract_comments))
            text = item["comment1"]  # 模型整体的输出，包括建议和内容
            # 修改后内容的正则
            modified_pattern = re.compile(r'修改后的内容：\n(.+)', re.S)
            modified_match = modified_pattern.search(text)
            modified_content = modified_match.group(1).strip() if modified_match else "未找到修改后的内容"
            item["comment1"] = modified_content
            print("建议提取结束")
        revise_document(revise_filepath, processed_filepath, contract_comments)
    except Exception as e:
        print("check_contract捕获到异常：", str(e))
        import traceback
        traceback.print_exc()  # 打印异常堆栈信息
        return None


def process_paragraph(para, index):
    api_url = url_base_list[index % 4]
    text = para.text.strip()
    # 删除两端空格，并将多个连续空格替换为一个空格
    text = re.sub(r'\s+', ' ', text)
    if len(text) > 20:
        print(text)
        try:
            advice = get_correction_suggestion(text)  # 增加超时机制
            print(advice)
            return para, index, advice
        except Exception as e:
            print(f"获取建议时发生异常：{e}")
            return None, index, None
    return None, index, None


def check_text(doc):
    total_paragraphs = len(doc.paragraphs)
    paragraphs_to_process = [{'para': para, 'index': i} for i, para in enumerate(doc.paragraphs) if para.text.strip()]
    removed_paragraphs_count = total_paragraphs - len(paragraphs_to_process)
    contract_comments = []

    # 创建一个进度条
    progress_bar = tqdm(total=len(paragraphs_to_process), desc="正文处理中", unit="段", position=0, leave=True)

    # 使用锁来同步多线程中的进度条更新
    progress_lock = threading.Lock()

    with ThreadPoolExecutor(max_workers=4) as executor:
        futures = [executor.submit(process_paragraph, item['para'], item['index']) for item in paragraphs_to_process]

        # 迭代等待每个任务完成
        for future in as_completed(futures):
            try:
                para, index, advice = future.result()
                print(advice)
                if advice:
                    print("doc.paragraphs[index].text", doc.paragraphs[index].text)
                    contract_comments.append({'text': doc.paragraphs[index].text, 'comment1': advice, 'index': index})
            except Exception as e:
                print(f"任务执行失败：{e}")

            # 使用锁来确保进度条的正确更新
            with progress_lock:
                progress_bar.update(1)

        # 关闭进度条
        progress_bar.close()

    return contract_comments, doc
